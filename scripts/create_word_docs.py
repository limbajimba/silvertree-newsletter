#!/usr/bin/env python3
"""
Convert documentation markdown files to professionally formatted Word documents.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


# SilverTree brand color
SILVERTREE_GREEN = RGBColor(0x1A, 0x5F, 0x2A)


def setup_styles(doc: Document):
    """Configure document styles for professional appearance."""
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = SILVERTREE_GREEN

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = SILVERTREE_GREEN

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = SILVERTREE_GREEN

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Normal text
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'

    return doc


def add_title_page(doc: Document, title: str, subtitle: str):
    """Add a professional title page."""
    # Add spacing at top
    for _ in range(6):
        doc.add_paragraph()

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("SILVERTREE EQUITY")
    run.font.size = Pt(16)
    run.font.color.rgb = SILVERTREE_GREEN
    run.font.bold = True

    doc.add_paragraph()

    # Document title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add spacing
    for _ in range(8):
        doc.add_paragraph()

    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run("Document Version: 1.0\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = info_para.add_run("Last Updated: January 2026\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = info_para.add_run("Classification: Internal Use Only")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break
    doc.add_page_break()


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse markdown table into headers and rows."""
    headers = []
    rows = []

    for i, line in enumerate(lines):
        if not line.strip() or line.strip().startswith('|-'):
            continue

        cells = [cell.strip() for cell in line.strip('|').split('|')]
        cells = [c.strip() for c in cells if c.strip() or len(cells) > 1]

        if not headers:
            headers = cells
        else:
            rows.append(cells)

    return headers, rows


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(header_cells):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data
                for para in row_cells[col_idx].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph()


def convert_markdown_to_docx(md_path: Path, docx_path: Path, title: str, subtitle: str):
    """Convert a markdown file to a Word document."""
    doc = Document()
    setup_styles(doc)

    # Add title page
    add_title_page(doc, title, subtitle)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Skip title page metadata (already added)
        if line.startswith('**Document Version') or line.startswith('**Last Updated') or line.startswith('**Classification'):
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_text = '\n'.join(code_lines)
                    run = code_para.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            headers, rows = parse_table(table_lines)
            add_table(doc, headers, rows)
            table_lines = []
            in_table = False

        # Headers
        if line.startswith('# ') and not line.startswith('## '):
            # Skip main title (already on title page)
            if 'SilverTree' not in line:
                text = line[2:].strip()
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Remove links
                doc.add_heading(text, level=0)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            # Remove anchor links
            text = re.sub(r'\{#[^}]+\}', '', text).strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Horizontal rule / section break
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # Skip TOC links
        if re.match(r'^\d+\.\s*\[', line.strip()):
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle bold
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            para = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            text = line.strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
            text = re.sub(r'`([^`]+)`', r'\1', text)  # Inline code
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links

            para = doc.add_paragraph(text)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        headers, rows = parse_table(table_lines)
        add_table(doc, headers, rows)

    # Save document
    doc.save(str(docx_path))
    print(f"Created: {docx_path}")


def main():
    docs_dir = Path(__file__).parent.parent / 'docs'

    # System Architecture
    convert_markdown_to_docx(
        docs_dir / 'SYSTEM_ARCHITECTURE.md',
        docs_dir / 'SilverTree_System_Architecture.docx',
        'Market Intelligence System',
        'System Architecture Documentation'
    )

    # Tracking Scope
    convert_markdown_to_docx(
        docs_dir / 'TRACKING_SCOPE.md',
        docs_dir / 'SilverTree_Tracking_Scope.docx',
        'Market Intelligence System',
        'Tracking Scope Documentation'
    )

    print("\nWord documents created in docs/ directory.")
    print("Open in Word and export as PDF via File > Export > Create PDF/XPS")


if __name__ == '__main__':
    main()
